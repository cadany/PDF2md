import os
import re
import uuid
import datetime
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shape import InlineShape
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging


from config import get_config


class DOCXService:
    def __init__(self, output_dir: str = None):
        """
        初始化DOCX服务
        
        Args:
            upload_dir: 上传文件存储目录，如果为None则从配置中读取
        """
        # 获取配置
        self.config = get_config()
        
        # 设置上传目录
        if output_dir is None:
            self.upload_dir = self.config.get('storage.upload_dir', 'uploads')
        else:
            self.upload_dir = output_dir
            
        self.logger = logging.getLogger("service.DOCXService")
        # 使用uvicorn的日志配置，不添加自定义处理器
        
        # 确保上传目录存在
        os.makedirs(self.upload_dir, exist_ok=True)
        
        self.logger.info(f"DOCXService初始化完成，上传目录: {self.upload_dir}")

    def extract_document_structure(self, docx_path: str) -> Dict[str, Any]:
        """
        提取Word文档结构（大纲/目录）
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            Dict: 文档结构信息
        """
        try:
            doc = Document(docx_path)
            structure = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_images': len(doc.inline_shapes),
                'sections': [],
                'hierarchy': [],
                'tables': [],
                'images': []
            }
            
            # 提取大纲结构
            current_level = 0
            current_section = None
            
            for i, paragraph in enumerate(doc.paragraphs):
                # 检查段落是否为标题
                heading_level = self._get_heading_level(paragraph)
                
                if heading_level > 0:
                    # 这是一个标题段落
                    if current_section:
                        # 设置前一个章节的结束位置（当前标题的前一个段落）
                        current_section['end_paragraph'] = i - 1
                        # 保存当前章节
                        structure['sections'].append(current_section)
                    
                    # 创建新章节
                    current_section = {
                        'title': paragraph.text.strip(),
                        'level': heading_level,
                        'start_paragraph': i,
                        'end_paragraph': None,
                        'tables': [],
                        'images': [],
                        'content_paragraphs': []
                    }
                    
                    # 更新层级结构
                    self._update_hierarchy(structure['hierarchy'], heading_level, paragraph.text.strip())
                    
                elif current_section:
                    # 这是正文内容，添加到当前章节
                    if paragraph.text.strip():
                        current_section['content_paragraphs'].append({
                            'text': paragraph.text.strip(),
                            'index': i
                        })
            
            # 处理最后一个章节
            if current_section:
                # 设置最后一个章节的结束位置（文档末尾）
                current_section['end_paragraph'] = len(doc.paragraphs) - 1
                structure['sections'].append(current_section)
            
            # 提取表格和图片信息
            self._extract_tables_and_images(doc, structure)
            
            self.logger.info(f"成功提取文档结构，共{len(structure['sections'])}个章节")
            return structure
            
        except Exception as e:
            self.logger.error(f"提取文档结构失败: {e}")
            raise

    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """
        获取段落标题级别
        
        Args:
            paragraph: 段落对象
            
        Returns:
            int: 标题级别（0表示不是标题）
        """
        if paragraph.style.name.startswith('Heading'):
            # 提取Heading后面的数字
            match = re.search(r'Heading (\d+)', paragraph.style.name)
            if match:
                return int(match.group(1))
        
        # 检查段落样式是否为标题样式
        if hasattr(paragraph.style, 'base_style') and paragraph.style.base_style:
            if paragraph.style.base_style.name.startswith('Heading'):
                match = re.search(r'Heading (\d+)', paragraph.style.base_style.name)
                if match:
                    return int(match.group(1))
        
        return 0

    def _update_hierarchy(self, hierarchy: List[Dict], level: int, title: str):
        """
        更新层级结构
        
        Args:
            hierarchy: 层级结构列表
            level: 当前级别
            title: 标题文本
        """
        # 找到当前级别的父级
        while hierarchy and hierarchy[-1]['level'] >= level:
            hierarchy.pop()
        
        # 添加当前级别
        current_node = {
            'level': level,
            'title': title,
            'children': []
        }
        
        if hierarchy:
            hierarchy[-1]['children'].append(current_node)
        else:
            hierarchy.append(current_node)

    def _extract_tables_and_images(self, doc: Document, structure: Dict[str, Any]):
        """
        提取表格和图片信息
        
        Args:
            doc: Word文档对象
            structure: 文档结构
        """
        # 提取表格信息
        for i, table in enumerate(doc.tables):
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'content': []
            }
            
            # 提取表格内容
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                table_info['content'].append(row_content)
            
            structure['tables'].append(table_info)
        
        # 提取图片信息
        for i, shape in enumerate(doc.inline_shapes):
            if shape.type == 3:  # 图片类型
                image_info = {
                    'index': i,
                    'width': shape.width,
                    'height': shape.height
                }
                structure['images'].append(image_info)

    def get_leaf_sections(self, structure: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        获取所有章节列表（每个标题都保存为单独文件）
        
        Args:
            structure: 文档结构
            
        Returns:
            List: 所有章节列表
        """
        all_sections = []
        
        for section in structure['sections']:
            # 为每个章节创建完整信息
            section_info = {
                'title': section['title'],
                'full_title': section['title'],  # 简化处理，不使用层级路径
                'level': section['level'],
                'content': section['content_paragraphs'],
                'start_paragraph': section['start_paragraph'],
                'end_paragraph': section['end_paragraph']
            }
            all_sections.append(section_info)
        
        self.logger.info(f"找到{len(all_sections)}个章节，将全部保存为单独文件")
        return all_sections

    def _get_section_content(self, structure: Dict[str, Any], section_title: str) -> List[Dict]:
        """
        获取指定章节的内容段落
        
        Args:
            structure: 文档结构
            section_title: 章节标题
            
        Returns:
            List: 内容段落列表
        """
        content = []
        
        for section in structure['sections']:
            if section['title'] == section_title:
                content = section['content_paragraphs']
                break
        
        return content

    def save_leaf_section_as_docx(self, docx_path: str, leaf_section: Dict[str, Any], 
                                 output_dir: str = None) -> str:
        """
        将末级目录内容保存为单独的DOCX文件
        
        Args:
            docx_path: 原始Word文档路径
            leaf_section: 末级章节信息
            output_dir: 输出目录，默认为上传目录
            
        Returns:
            str: 保存的文件路径
        """
        if output_dir is None:
            output_dir = self.upload_dir
        
        # 生成文件名：末级目录名称_年时分秒毫秒.docx
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")[:-3]  # 去掉最后3位微秒
        safe_title = self._sanitize_filename(leaf_section['title'])
        filename = f"{safe_title}_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)
        
        try:
            # 打开原始文档
            source_doc = Document(docx_path)
            
            # 创建新文档
            new_doc = Document()
            
            # 复制文档属性
            self._copy_document_properties(source_doc, new_doc)
            
            # 添加章节标题
            title_paragraph = new_doc.add_heading(leaf_section['full_title'], level=1)
            
            # 复制内容段落
            self._copy_section_content(source_doc, new_doc, leaf_section)
            
            # 保存新文档
            new_doc.save(output_path)
            
            self.logger.info(f"末级目录内容已保存到: {output_path}")
            return output_path
        except Exception as e:
            self.logger.error(f"保存末级目录内容失败: {e}")
            raise

    def _sanitize_filename(self, filename: str) -> str:
        """
        清理文件名，移除非法字符
        
        Args:
            filename: 原始文件名
            
        Returns:
            str: 清理后的文件名
        """
        # 移除非法字符
        invalid_chars = r'[<>:"/\\|?*]'
        safe_name = re.sub(invalid_chars, '_', filename)
        
        # 限制文件名长度
        if len(safe_name) > 100:
            safe_name = safe_name[:100]
        
        return safe_name

    def _copy_document_properties(self, source_doc: Document, target_doc: Document):
        """
        复制文档属性
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
        """
        # 复制核心属性
        core_props = source_doc.core_properties
        target_core_props = target_doc.core_properties
        
        target_core_props.title = core_props.title or ""
        target_core_props.subject = core_props.subject or ""
        target_core_props.author = core_props.author or ""
        target_core_props.keywords = core_props.keywords or ""
        target_core_props.comments = core_props.comments or ""

    def _copy_section_content(self, source_doc: Document, target_doc: Document, 
                            leaf_section: Dict[str, Any]):
        """
        深度复制章节内容，保持格式和顺序
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            leaf_section: 末级章节信息
        """
        # 复制章节标题（使用原始格式）
        start_idx = leaf_section['start_paragraph']
        end_idx = leaf_section['end_paragraph'] if leaf_section['end_paragraph'] else len(source_doc.paragraphs)
        
        # 深度复制：保持文档元素的原始顺序
        self._deep_copy_content(source_doc, target_doc, start_idx, end_idx)

    def _deep_copy_content(self, source_doc: Document, target_doc: Document, 
                          start_idx: int, end_idx: int):
        """
        深度复制文档内容，保持格式和顺序
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 简化实现：直接按段落索引复制，同时处理表格
        self._simple_deep_copy(source_doc, target_doc, start_idx, end_idx)

    def _simple_deep_copy(self, source_doc: Document, target_doc: Document, 
                         start_idx: int, end_idx: int):
        """
        简化但可靠的深度复制实现
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 复制段落
        for i in range(start_idx, end_idx):
            if i < len(source_doc.paragraphs):
                source_paragraph = source_doc.paragraphs[i]
                
                # 检查是否为标题段落
                heading_level = self._get_heading_level(source_paragraph)
                
                if heading_level > 0:
                    # 添加标题
                    new_heading = target_doc.add_heading(source_paragraph.text.strip(), level=heading_level)
                    # 深度复制标题格式
                    self._deep_copy_paragraph_format(source_paragraph, new_heading)
                else:
                    # 添加普通段落
                    new_paragraph = target_doc.add_paragraph()
                    # 深度复制段落格式
                    self._deep_copy_paragraph_format(source_paragraph, new_paragraph)
        
        # 复制表格（简化方法：复制所有表格）
        self._copy_all_tables(source_doc, target_doc)

    def _copy_all_tables(self, source_doc: Document, target_doc: Document):
        """
        复制所有表格（简化实现）
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
        """
        for source_table in source_doc.tables:
            try:
                # 创建新表格
                rows = len(source_table.rows)
                cols = len(source_table.columns)
                new_table = target_doc.add_table(rows=rows, cols=cols)
                
                # 复制表格样式
                new_table.style = source_table.style
                
                # 深度复制表格内容
                for i, row in enumerate(source_table.rows):
                    for j, cell in enumerate(row.cells):
                        new_cell = new_table.cell(i, j)
                        
                        # 复制单元格文本
                        new_cell.text = cell.text
                        
                        # 深度复制单元格格式
                        if hasattr(cell, 'paragraphs') and cell.paragraphs:
                            for para in cell.paragraphs:
                                new_para = new_cell.add_paragraph()
                                self._deep_copy_paragraph_format(para, new_para)
                
                self.logger.info(f"成功复制表格，行数: {rows}, 列数: {cols}")
            except Exception as e:
                self.logger.warning(f"复制表格失败: {e}")

    def _get_elements_in_range(self, source_doc: Document, start_idx: int, end_idx: int) -> List[Dict]:
        """
        获取指定范围内的所有文档元素（段落、表格、图片）
        
        Args:
            source_doc: 源文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
            
        Returns:
            List: 元素信息列表
        """
        elements = []
        
        # 获取文档主体
        body = source_doc._element.body
        
        # 遍历文档中的所有元素
        for element in body:
            element_info = self._analyze_element(element, source_doc)
            if element_info and self._is_element_in_range(element_info, start_idx, end_idx):
                elements.append(element_info)
        
        # 按原始顺序排序
        elements.sort(key=lambda x: x.get('position', 0))
        
        return elements

    def _analyze_element(self, element, source_doc: Document) -> Optional[Dict]:
        """
        分析文档元素类型和位置
        
        Args:
            element: XML元素
            source_doc: 源文档
            
        Returns:
            Dict: 元素信息
        """
        try:
            # 检查元素类型
            tag = element.tag
            
            if 'p' in tag:  # 段落
                # 查找对应的段落对象
                for i, para in enumerate(source_doc.paragraphs):
                    if para._element == element:
                        return {
                            'type': 'paragraph',
                            'element': element,
                            'paragraph': para,
                            'position': i,
                            'index': i
                        }
            
            elif 'tbl' in tag:  # 表格
                # 查找对应的表格对象
                for i, table in enumerate(source_doc.tables):
                    if table._element == element:
                        return {
                            'type': 'table',
                            'element': element,
                            'table': table,
                            'position': i,
                            'index': i
                        }
            
            elif 'drawing' in tag:  # 图片
                return {
                    'type': 'image',
                    'element': element,
                    'position': len(source_doc.paragraphs) + len(source_doc.tables),
                    'index': len(source_doc.inline_shapes)
                }
        
        except Exception as e:
            self.logger.warning(f"分析元素失败: {e}")
        
        return None

    def _is_element_in_range(self, element_info: Dict, start_idx: int, end_idx: int) -> bool:
        """
        检查元素是否在指定范围内
        
        Args:
            element_info: 元素信息
            start_idx: 起始索引
            end_idx: 结束索引
            
        Returns:
            bool: 是否在范围内
        """
        if element_info['type'] == 'paragraph':
            return start_idx <= element_info['index'] <= end_idx
        elif element_info['type'] == 'table':
            # 表格的位置判断需要更复杂的逻辑
            # 暂时假设所有表格都在范围内
            return True
        elif element_info['type'] == 'image':
            # 图片的位置判断需要更复杂的逻辑
            # 暂时假设所有图片都在范围内
            return True
        
        return False

    def _deep_copy_paragraph(self, source_doc: Document, target_doc: Document, element_info: Dict):
        """
        深度复制段落
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            element_info: 段落元素信息
        """
        source_paragraph = element_info['paragraph']
        
        # 检查是否为标题段落
        heading_level = self._get_heading_level(source_paragraph)
        
        if heading_level > 0:
            # 添加标题
            new_heading = target_doc.add_heading(source_paragraph.text.strip(), level=heading_level)
            # 深度复制标题格式
            self._deep_copy_paragraph_format(source_paragraph, new_heading)
        else:
            # 添加普通段落
            new_paragraph = target_doc.add_paragraph()
            # 深度复制段落格式
            self._deep_copy_paragraph_format(source_paragraph, new_paragraph)

    def _deep_copy_paragraph_format(self, source_paragraph: Paragraph, target_paragraph: Paragraph):
        """
        深度复制段落格式
        
        Args:
            source_paragraph: 源段落
            target_paragraph: 目标段落
        """
        try:
            # 复制段落样式
            if source_paragraph.style:
                target_paragraph.style = source_paragraph.style
            
            # 复制段落对齐方式
            if source_paragraph.alignment:
                target_paragraph.alignment = source_paragraph.alignment
            
            # 复制段落格式属性
            source_format = source_paragraph.paragraph_format
            target_format = target_paragraph.paragraph_format
            
            # 复制缩进
            if hasattr(source_format, 'left_indent'):
                target_format.left_indent = source_format.left_indent
            if hasattr(source_format, 'right_indent'):
                target_format.right_indent = source_format.right_indent
            if hasattr(source_format, 'first_line_indent'):
                target_format.first_line_indent = source_format.first_line_indent
            
            # 复制间距
            if hasattr(source_format, 'space_before'):
                target_format.space_before = source_format.space_before
            if hasattr(source_format, 'space_after'):
                target_format.space_after = source_format.space_after
            if hasattr(source_format, 'line_spacing'):
                target_format.line_spacing = source_format.line_spacing
            
            # 深度复制文本运行
            for source_run in source_paragraph.runs:
                new_run = target_paragraph.add_run(source_run.text)
                self._deep_copy_run_format(source_run, new_run)
                
        except Exception as e:
            self.logger.warning(f"深度复制段落格式失败: {e}")

    def _deep_copy_run_format(self, source_run, target_run):
        """
        深度复制文本运行格式
        
        Args:
            source_run: 源文本运行
            target_run: 目标文本运行
        """
        try:
            # 复制字体名称
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            
            # 复制字体大小
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            
            # 复制字体样式
            target_run.font.bold = source_run.font.bold
            target_run.font.italic = source_run.font.italic
            target_run.font.underline = source_run.font.underline
            
            # 复制字体颜色
            if hasattr(source_run.font, 'color') and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            
            # 复制高亮颜色
            if hasattr(source_run.font, 'highlight_color'):
                target_run.font.highlight_color = source_run.font.highlight_color
            
            # 复制字符间距
            if hasattr(source_run, 'font') and hasattr(source_run.font, 'spacing'):
                target_run.font.spacing = source_run.font.spacing
            
            # 复制其他字体属性
            if hasattr(source_run.font, 'strike'):
                target_run.font.strike = source_run.font.strike
            if hasattr(source_run.font, 'double_strike'):
                target_run.font.double_strike = source_run.font.double_strike
            if hasattr(source_run.font, 'shadow'):
                target_run.font.shadow = source_run.font.shadow
            if hasattr(source_run.font, 'outline'):
                target_run.font.outline = source_run.font.outline
            if hasattr(source_run.font, 'emboss'):
                target_run.font.emboss = source_run.font.emboss
            if hasattr(source_run.font, 'imprint'):
                target_run.font.imprint = source_run.font.imprint
                
        except Exception as e:
            self.logger.warning(f"深度复制文本格式失败: {e}")

    def _deep_copy_table(self, source_doc: Document, target_doc: Document, element_info: Dict):
        """
        深度复制表格
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            element_info: 表格元素信息
        """
        try:
            source_table = element_info['table']
            
            # 创建新表格
            rows = len(source_table.rows)
            cols = len(source_table.columns)
            new_table = target_doc.add_table(rows=rows, cols=cols)
            
            # 复制表格样式
            new_table.style = source_table.style
            
            # 深度复制表格内容
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    
                    # 复制单元格文本
                    new_cell.text = cell.text
                    
                    # 深度复制单元格格式
                    if hasattr(cell, 'paragraphs') and cell.paragraphs:
                        for para in cell.paragraphs:
                            new_para = new_cell.add_paragraph()
                            self._deep_copy_paragraph_format(para, new_para)
            
        except Exception as e:
            self.logger.warning(f"深度复制表格失败: {e}")

    def _deep_copy_image(self, source_doc: Document, target_doc: Document, element_info: Dict):
        """
        深度复制图片
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            element_info: 图片元素信息
        """
        # 图片复制需要更复杂的实现
        self.logger.info("检测到图片，图片复制功能需要更复杂的实现")
        
        # 暂时添加一个占位符段落
        target_doc.add_paragraph("[图片位置]")

    def _copy_tables_in_range(self, source_doc: Document, target_doc: Document, 
                            start_idx: int, end_idx: int):
        """
        复制指定段落范围内的表格
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 这个方法现在被深度复制方法替代
        pass
    
    def _copy_images_in_range(self, source_doc: Document, target_doc: Document, 
                            start_idx: int, end_idx: int):
        """
        复制指定段落范围内的图片
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 这个方法现在被深度复制方法替代
        pass

    def _copy_paragraph_format(self, source_paragraph: Paragraph, target_paragraph: Paragraph):
        """
        复制段落格式
        
        Args:
            source_paragraph: 源段落
            target_paragraph: 目标段落
        """
        try:
            # 复制段落对齐方式
            if source_paragraph.alignment:
                target_paragraph.alignment = source_paragraph.alignment
            
            # 复制段落样式
            if source_paragraph.style:
                target_paragraph.style = source_paragraph.style
            
            # 复制段落缩进
            if hasattr(source_paragraph.paragraph_format, 'left_indent'):
                target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
            
            if hasattr(source_paragraph.paragraph_format, 'right_indent'):
                target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
            
            if hasattr(source_paragraph.paragraph_format, 'first_line_indent'):
                target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
            
            # 复制行间距
            if hasattr(source_paragraph.paragraph_format, 'line_spacing'):
                target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
            
            # 复制段前段后间距
            if hasattr(source_paragraph.paragraph_format, 'space_before'):
                target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
            
            if hasattr(source_paragraph.paragraph_format, 'space_after'):
                target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
                
        except Exception as e:
            self.logger.warning(f"复制段落格式失败: {e}")

    def _copy_run_format(self, source_run, target_run):
        """
        复制文本运行格式
        
        Args:
            source_run: 源文本运行
            target_run: 目标文本运行
        """
        try:
            # 复制字体名称
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            
            # 复制字体大小
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            
            # 复制字体样式
            target_run.font.bold = source_run.font.bold
            target_run.font.italic = source_run.font.italic
            target_run.font.underline = source_run.font.underline
            
            # 复制字体颜色
            if hasattr(source_run.font, 'color') and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            
            # 复制高亮颜色
            if hasattr(source_run.font, 'highlight_color'):
                target_run.font.highlight_color = source_run.font.highlight_color
            
            # 复制字符间距
            if hasattr(source_run, 'font') and hasattr(source_run.font, 'spacing'):
                target_run.font.spacing = source_run.font.spacing
                
        except Exception as e:
            self.logger.warning(f"复制文本格式失败: {e}")

    def process_document(self, docx_path: str, output_dir: str = None) -> Dict[str, Any]:
        """
        处理Word文档，提取末级目录并保存为单独文件
        
        Args:
            docx_path: Word文档路径
            output_dir: 输出目录
            
        Returns:
            Dict: 处理结果
        """
        try:
            # 提取文档结构
            structure = self.extract_document_structure(docx_path)
            # 获取末级目录
            leaf_sections = self.get_leaf_sections(structure)
            
            # 保存每个末级目录为单独文件
            saved_files = []
            for leaf_section in leaf_sections:
                output_path = self.save_leaf_section_as_docx(docx_path, leaf_section, output_dir)
                saved_files.append({
                    'title': leaf_section['title'],
                    'full_title': leaf_section['full_title'],
                    'output_path': output_path,
                    'filename': os.path.basename(output_path)
                })
            
            result = {
                'success': True,
                'original_file': docx_path,
                'total_sections': len(leaf_sections),
                'saved_files': saved_files,
                'structure_summary': {
                    'total_paragraphs': structure['total_paragraphs'],
                    'total_tables': structure['total_tables'],
                    'total_images': structure['total_images']
                }
            }
            
            self.logger.info(f"文档处理完成，共保存{len(saved_files)}个末级目录文件")
            return result
            
        except Exception as e:
            self.logger.error(f"文档处理失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'original_file': docx_path
            }

    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        获取DOCX文件信息
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 文件信息或None
        """
        if not os.path.exists(file_path):
            return None
        
        try:
            doc = Document(file_path)
            
            return {
                'file_path': file_path,
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_images': len(doc.inline_shapes),
                'core_properties': {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None
                }
            }
        except Exception as e:
            self.logger.error(f"获取文件信息失败: {e}")
            return None

    def _copy_tables_in_range(self, source_doc: Document, target_doc: Document, 
                            start_idx: int, end_idx: int):
        """
        复制指定段落范围内的表格
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 由于python-docx没有直接的方法获取表格所在的段落位置，
        # 我们采用遍历所有表格并检查其相对位置的方法
        for table in source_doc.tables:
            # 获取表格在文档中的大致位置
            # 这里使用一个简单的启发式方法：找到表格前后的段落
            table_start_idx = self._find_table_start_position(source_doc, table)
            
            if table_start_idx is not None and start_idx <= table_start_idx <= end_idx:
                # 表格在目标范围内，复制表格
                self._copy_table(table, target_doc)
    
    def _find_table_start_position(self, doc: Document, table) -> Optional[int]:
        """
        查找表格的起始位置（段落索引）
        
        Args:
            doc: 文档对象
            table: 表格对象
            
        Returns:
            int: 表格起始段落索引，如果找不到返回None
        """
        # 由于python-docx的API限制，无法直接获取表格所在的精确段落位置
        # 这里使用一个简化的方法：遍历所有段落，找到表格前的标题段落
        
        # 获取表格在文档元素中的位置
        table_element = table._element
        
        # 遍历文档中的所有段落元素
        for i, paragraph in enumerate(doc.paragraphs):
            para_element = paragraph._element
            
            # 检查段落元素是否在表格元素之前
            # 这是一个简化的位置比较
            if para_element is not None and table_element is not None:
                # 如果段落包含对表格的引用，或者段落紧邻表格，则认为是表格的起始位置
                if para_element.getparent() == table_element.getparent():
                    # 检查段落和表格是否在同一父级下
                    parent = para_element.getparent()
                    if parent is not None:
                        # 获取段落和表格在父级中的索引
                        para_index = parent.index(para_element)
                        table_index = parent.index(table_element)
                        
                        # 如果段落紧邻表格之前，则认为是表格的起始位置
                        if para_index + 1 == table_index:
                            return i
        
        # 如果找不到精确位置，返回表格大致位置（使用第一个段落作为参考）
        # 这是一个fallback方法
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text and '表' in paragraph.text:
                # 如果段落包含"表"字，可能是表格标题
                return i
        
        return None
    
    def _copy_table(self, source_table, target_doc: Document):
        """
        复制表格
        
        Args:
            source_table: 源表格
            target_doc: 目标文档
        """
        # 创建新表格
        rows = len(source_table.rows)
        cols = len(source_table.columns)
        new_table = target_doc.add_table(rows=rows, cols=cols)
        
        # 复制表格样式
        new_table.style = source_table.style
        
        # 复制表格内容
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                
                # 复制单元格文本
                new_cell.text = cell.text
                
                # 复制单元格格式
                if hasattr(cell, 'paragraphs') and cell.paragraphs:
                    for para in cell.paragraphs:
                        new_para = new_cell.add_paragraph()
                        self._copy_paragraph_format(para, new_para)
                        
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            self._copy_run_format(run, new_run)
    
    def _copy_images_in_range(self, source_doc: Document, target_doc: Document, 
                            start_idx: int, end_idx: int):
        """
        复制指定段落范围内的图片
        
        Args:
            source_doc: 源文档
            target_doc: 目标文档
            start_idx: 起始段落索引
            end_idx: 结束段落索引
        """
        # 复制内联图片
        for shape in source_doc.inline_shapes:
            # 检查图片是否在目标段落范围内
            if hasattr(shape, '_inline') and hasattr(shape._inline, 'graphic'):
                # 这是一个简化的实现，实际图片复制需要更复杂的处理
                # 这里暂时记录日志，表示检测到图片
                self.logger.info(f"检测到图片，但图片复制功能需要更复杂的实现")

def main():
    docx_service = DOCXService(output_dir='files/output')
    file_info = docx_service.process_document('./files/fj2.docx')
    # print(file_info)

if __name__ == '__main__':
    main()
